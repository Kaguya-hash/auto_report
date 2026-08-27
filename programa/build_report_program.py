"""
Builds the scoring report as a Word document.

This file has NO domain knowledge: it doesn't know what ADI-R is, which
items exist, how scores are combined, or what the cut-offs are. All of
that lives in the encrypted configuration loaded through
secure_config.py and is executed by the generic rule engine in
engine.py. Someone reading this file only learns that a report is built
from a schema of sections/subsections, each with a "rule" evaluated
against the person's answers -- nothing about the actual algorithm.
"""

from pypdf import PdfReader

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .secure_config import load_config
from .engine import evaluate

from pathlib import Path

COLUMN_SIZES = [85, 15]
PDF_PATH = Path(__file__).resolve().parent.parent / "questions_now.pdf"
OUTPUT_PATH = Path(__file__).resolve().parent.parent / "rapport.docx"


def set_column_percentages(table, percentages):
    """
    Set table columns as percentages of the full available table width.
    The table itself occupies 100% of the available width.
    """
    if len(percentages) != len(table.columns):
        raise ValueError("Number of percentages must match number of columns.")

    if abs(sum(percentages) - 100) > 1e-6:
        raise ValueError("Percentages must sum to 100.")

    table.autofit = False

    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)

    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    for column, percentage in zip(table.columns, percentages):
        width = int(5000 * percentage / 100)
        for cell in column.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "pct")


def read_pdf_fields(path):
    reader = PdfReader(path)
    fields_brute = reader.get_fields()
    return {name: info.get("/V", "No Answer") for name, info in fields_brute.items()}


def compute_variables(fields, config):
    """Compute the small set of derived variables (e.g. age) the rules rely on."""
    variables = {}
    for name, spec in config.get("derived_variables", {}).items():
        total = 0.0
        for term in spec["terms"]:
            value = float(fields[term["field"]])
            if "divide" in term:
                value /= term["divide"]
            total += value
        variables[name] = total
    return variables


def add_new_row(table, first_row_flag):
    return table.rows[0].cells if first_row_flag else table.add_row().cells


def build_section(doc, section_key, section_cfg, fields, variables, rel_items, value_maps):
    doc.add_heading(f"{section_key}: {section_cfg['title']}\n", level=3)

    subsections = section_cfg.get("subsections")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_column_percentages(table, COLUMN_SIZES)

    if subsections:
        computed = {}
        first_row = True

        for sub_key, sub_cfg in subsections.items():
            row_cells = add_new_row(table, first_row)
            first_row = False

            row_cells[0].text = f"{sub_key}: {sub_cfg['table_title']}"

            try:
                value, sentence = evaluate(
                    sub_cfg["value_rule"], fields, variables, rel_items, value_maps, computed
                )
                computed[sub_key] = (value, sentence)
                row_cells[1].text = str(value)
            except Exception:
                computed[sub_key] = (0, "")
                row_cells[1].text = "?"

        total_row = table.add_row().cells
        total_row[0].text = "Résultat: "

        try:
            total_value, _ = evaluate(
                section_cfg["total_rule"], fields, variables, rel_items, value_maps, computed
            )
        except:
            total_value = "?"

        try:
            cutoff_value, _ = evaluate(
                section_cfg["cutoff_rule"], fields, variables, rel_items, value_maps, computed
            )
        except:
            cutoff_value = "?"

        total_row[1].text = f"{total_value} / {cutoff_value}"

        total_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_row[0].paragraphs[0].runs[0].bold = True
        total_row[1].paragraphs[0].runs[0].bold = True

        for sub_key, sub_cfg in subsections.items():
            doc.add_heading(f"{sub_key}: {sub_cfg['paragraph_title']}\n", level=5)
            _, sentence = computed.get(sub_key, (0, ""))
            doc.add_paragraph(sentence if sentence else "?")

    else:
        row_cells = table.rows[0].cells
        row_cells[0].text = f"{section_key}: {section_cfg['title']}"

        try:
            value, sentence = evaluate(section_cfg["value_rule"], fields, variables, rel_items, value_maps, {})
        except:
            value, sentence = "?", "?"

        try:
            cutoff_value, _ = evaluate(section_cfg["cutoff_rule"], fields, variables, rel_items, value_maps, {})
        except:
            cutoff_value = "?"

        row_cells[1].text = f"{value} / {cutoff_value}"

        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].paragraphs[0].runs[0].bold = True

        doc.add_paragraph("\n" + (sentence or "?"))


def build_report(pdf_path: Path, output_path: Path):
    config = load_config(
        use_env_file = True
    )
    fields = read_pdf_fields(pdf_path)
    variables = compute_variables(fields, config)
    value_maps = {"cot_to_score": config["cot_to_score"]}

    doc = docx.Document()
    doc.add_heading(config["title"] + "\n", level=1)
    doc.add_paragraph(config["generic_text"])

    for section_key, section_cfg in config["sections"].items():
        build_section(doc, section_key, section_cfg, fields, variables, config["rel_items"], value_maps)

    doc.save(output_path)
